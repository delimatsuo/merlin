"""DOCX document generation service."""

import io
import re

import structlog
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = structlog.get_logger()


def _add_formatted_text(paragraph, text: str):
    """Add text to a paragraph, handling **bold** markdown syntax."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_thin_border(paragraph):
    """Add a thin bottom border to a paragraph (for section headings)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


async def generate_resume_docx(resume_content: str) -> bytes:
    """Generate a professional resume DOCX from markdown content."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    lines = resume_content.strip().split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line:
            continue

        # H1 — Candidate name
        if line.startswith("# "):
            name_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_after = Pt(2)
            continue

        # H2 — Section headings
        if line.startswith("## "):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text.upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            _add_thin_border(p)
            continue

        # H3 — Experience/education entries
        if line.startswith("### "):
            entry_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _add_formatted_text(p, entry_text)
            for run in p.runs:
                run.font.size = Pt(10)
                if not run.bold:
                    run.bold = True
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            _add_formatted_text(p, bullet_text)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Cm(0.5)
            continue

        # Regular text (contact info, dates, skills, etc.)
        p = doc.add_paragraph()
        # Contact info line (contains | or @ typically right after name)
        if "|" in line and ("@" in line or "(" in line):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_text(p, line)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(6)
        else:
            _add_formatted_text(p, line)
            for run in p.runs:
                run.font.size = Pt(10)
            # Date lines (e.g., "01/2025 – Atual") — make them smaller/gray
            if re.match(r'^\d{2}/\d{4}\s*[–\-—]', line):
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.paragraph_format.space_after = Pt(2)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


async def generate_cover_letter_docx(cover_letter_text: str) -> bytes:
    """Generate a cover letter DOCX."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    paragraphs = cover_letter_text.strip().split("\n\n")

    for i, para_text in enumerate(paragraphs):
        para_text = para_text.strip()
        if not para_text:
            continue

        p = doc.add_paragraph()
        _add_formatted_text(p, para_text)

        if i == 0:
            p.paragraph_format.space_after = Pt(12)
        else:
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
