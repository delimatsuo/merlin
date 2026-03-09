"""Resume file parsing service (PDF + DOCX)."""

import io

import structlog

logger = structlog.get_logger()


async def parse_resume(content: bytes, content_type: str) -> str:
    """Parse resume file content to plain text."""
    if content_type == "application/pdf":
        return _parse_pdf(content)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(content)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n\n".join(text_parts)


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx with defusedxml."""
    # Patch xml parsing for security before importing docx
    import defusedxml.minidom
    import xml.dom.minidom
    xml.dom.minidom.parseString = defusedxml.minidom.parseString

    from docx import Document

    doc = Document(io.BytesIO(content))
    text_parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)
